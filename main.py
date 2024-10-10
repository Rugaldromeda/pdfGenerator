import docx
from docx.enum.text import WD_BREAK
from docx.oxml.ns import nsdecls, qn
import os
import reportlab
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import ParagraphStyle
from findSize import findSize
from findBreak import findBreak


def createPDF():
    choiceDoc = input("Insira o endereço do arquivo: ").replace('"',"")
    name = input("Informe o nome do arquivo: ")
    index = int(input("Informe a numeração inicial do arquivo: "))
    documento = docx.Document(choiceDoc)
    nomeado = '{}{}.pdf'.format(name,index)

    separador = r'\\'

    origemDoc = os.path.dirname(choiceDoc.replace(separador,"/"))

    texto =[]

    for i in range(len(documento.paragraphs)):
        paragrafo = documento.paragraphs[i].text
        fontSize = 12

        if findSize(documento,i) != None:
            fontSize = findSize(documento,i)
        
        if paragrafo.strip():
            
            estiloPar = ParagraphStyle(
                'Normal',
                fontName= 'Helvetica',
                fontSize= fontSize,
                leading= fontSize * 1.15
            )
            
            texto.append(Paragraph(paragrafo, estiloPar))
        else:
            texto.append(Spacer(1,12))
        
        if findBreak(documento,i):
            pdf = SimpleDocTemplate('{}\\{}{}.pdf'.format(origemDoc,name,index), pagesize=A4)
            pdf.build(texto)
            texto.clear()
            index +=1
    
    pdf = SimpleDocTemplate('{}\\{}{}.pdf'.format(origemDoc,name,index), pagesize=A4)
    pdf.build(texto)

createPDF()